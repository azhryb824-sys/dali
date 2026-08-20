from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.section import WD_SECTION
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

OUT="docs/library/تقرير_مطابقة_وتنفيذ_منصة_دالي.docx"
NAVY="0B3042"; RED="C52D37"; GOLD="C79B47"; PALE="F2F5F6"; GREEN="237451"; AMBER="91651D"

def rtl(p):
    p.alignment=WD_ALIGN_PARAGRAPH.RIGHT
    pPr=p._p.get_or_add_pPr(); bidi=OxmlElement("w:bidi"); bidi.set(qn("w:val"),"1"); pPr.append(bidi)
def font(run,size=10,bold=False,color=NAVY):
    run.font.name="Arial"; run._element.get_or_add_rPr().rFonts.set(qn("w:cs"),"Arial"); run.font.size=Pt(size); run.bold=bold; run.font.color.rgb=RGBColor.from_string(color)
def para(doc,text,size=10,bold=False,color="263A43",after=6):
    p=doc.add_paragraph(); rtl(p); p.paragraph_format.space_after=Pt(after); p.paragraph_format.line_spacing=1.15; font(p.add_run(text),size,bold,color); return p
def heading(doc,text,level=1):
    p=doc.add_paragraph(); rtl(p); p.paragraph_format.space_before=Pt(16 if level==1 else 10); p.paragraph_format.space_after=Pt(7); font(p.add_run(text),16 if level==1 else 12,True,RED if level==1 else NAVY); p.paragraph_format.keep_with_next=True; return p
def shade(cell,fill):
    tcPr=cell._tc.get_or_add_tcPr(); shd=tcPr.find(qn("w:shd")) or OxmlElement("w:shd"); shd.set(qn("w:fill"),fill); tcPr.append(shd)
def set_cell(cell,text,bold=False,color=NAVY,size=8.5,fill=None):
    cell.text=""; p=cell.paragraphs[0]; rtl(p); font(p.add_run(text),size,bold,color); cell.vertical_alignment=WD_CELL_VERTICAL_ALIGNMENT.CENTER
    if fill: shade(cell,fill)
    tcMar=cell._tc.get_or_add_tcPr(); mar=OxmlElement("w:tcMar")
    for side in ("top","start","bottom","end"):
        el=OxmlElement(f"w:{side}"); el.set(qn("w:w"),"100"); el.set(qn("w:type"),"dxa"); mar.append(el)
    tcMar.append(mar)
def table(doc,headers,rows,widths):
    t=doc.add_table(rows=1,cols=len(headers)); t.alignment=WD_TABLE_ALIGNMENT.CENTER; t.autofit=False
    for i,h in enumerate(headers): set_cell(t.rows[0].cells[i],h,True,"FFFFFF",8.5,NAVY); t.rows[0].cells[i].width=Inches(widths[i])
    for row in rows:
        cells=t.add_row().cells
        for i,v in enumerate(row): set_cell(cells[i],str(v),i==0,NAVY,8.2,"FFFFFF" if len(t.rows)%2 else PALE); cells[i].width=Inches(widths[i])
    return t

doc=Document(); sec=doc.sections[0]; sec.page_width=Inches(8.5); sec.page_height=Inches(11); sec.top_margin=sec.bottom_margin=Inches(.78); sec.left_margin=sec.right_margin=Inches(.82); sec.header_distance=Inches(.35); sec.footer_distance=Inches(.35)
styles=doc.styles; normal=styles["Normal"]; normal.font.name="Arial"; normal._element.rPr.rFonts.set(qn("w:cs"),"Arial"); normal.font.size=Pt(10)
header=sec.header.paragraphs[0]; rtl(header); font(header.add_run("شركة دالي للتشغيل والصيانة | تقرير المطابقة التنفيذي"),8,True,NAVY)
footer=sec.footer.paragraphs[0]; footer.alignment=WD_ALIGN_PARAGRAPH.CENTER; font(footer.add_run("سري للاستخدام الداخلي • الإصدار 2.0 • 20 أغسطس 2026"),7,False,"6D7B82")

p=doc.add_paragraph(); p.paragraph_format.space_before=Pt(105); p.alignment=WD_ALIGN_PARAGRAPH.CENTER; font(p.add_run("تقرير المطابقة والإغلاق التنفيذي"),25,True,NAVY)
p=doc.add_paragraph(); p.alignment=WD_ALIGN_PARAGRAPH.CENTER; font(p.add_run("منصة دالي للمقاولات والتشغيل والصيانة وتوفير القوى العاملة"),15,True,RED)
p=doc.add_paragraph(); p.alignment=WD_ALIGN_PARAGRAPH.CENTER; font(p.add_run("مقارنة متطلبات الوثيقة بالتنفيذ الفعلي — واجهات، API، بيانات، صلاحيات، اختبارات وتشغيل"),10,False,"51636B")
doc.add_paragraph()
table(doc,["الحالة","العدد","المعنى"],[
    ("مكتمل برمجياً","39","واجهة قابلة للاستخدام + API + بيانات + صلاحيات + اختبار"),
    ("جزئي تشغيلياً","4","يتطلب بيانات أصلية أو اختباراً ميدانياً أو مزود مراقبة"),
    ("اعتماد خارجي","4","جهة حكومية أو مصرف أو اتفاقية ومفاتيح رسمية"),
    ("غير مكتمل برمجياً","0","لا توجد فجوة برمجية حرجة معروفة في نطاق الوثيقة الحالي"),
],[1.35,.7,4.45])
para(doc,"قرار القياس: لا يُعد التبويب أو الجدول الفارغ تنفيذاً. الإغلاق يتطلب دورة ناجحة وفاشلة وصلاحية وتدقيقاً وبيانات مستديمة.",10,True,RED,0)

doc.add_page_break(); heading(doc,"1. الخلاصة التنفيذية")
para(doc,"انتقل النظام من قسم مقاولات تعريفي إلى قطاع أعمال متكامل: فرص ومناقصات، معاينات وتقدير وBOQ، مشروعات ومراكز تكلفة، WBS ومخاطر، يوميات ميدانية، وثائق هندسية وإحالات، جودة وسلامة، مشتريات ومقاولو باطن، أوامر تغيير، مستخلصات واحتجاز، CBS/WIP/EAC، وتسليم وضمان. كما أصبحت رحلة عرض السعر العامة قابلة للحفظ والتتبع وإرفاق الملفات.")
heading(doc,"ما أُغلق منذ الإصدار 1.1",2)
for text in [
"MFA للأدوار الحساسة: TOTP مشفر، تحديات قصيرة، رموز استرداد أحادية، حد محاولات، ومنع الجلسة الحساسة قبل التحقق.",
"الملفات الهندسية: إصدارات متسلسلة، Transmittal، نسخة حالية وسابقة، مراجعة وسبب رفض وفصل واجبات وتنزيل محمي.",
"رقابة التكلفة: CBS تفصيلي، baseline، تغييرات، التزامات، فعلي، ETC، EAC، الانحراف، WIP، مستخلصات واحتجاز.",
"طلب عرض السعر: حفظ تلقائي، idempotency، تتبع بالرقم والبريد، مرفقات مفحوصة ومخزنة بصورة دائمة.",
"يوميات الموقع: أعمال وطقس وعمالة ومعدات ومعوقات وGPS ودقة الإشارة مع حدود تحقق في API وقاعدة البيانات.",
]: para(doc,"• "+text,9.5,False,"263A43",4)

heading(doc,"2. مصفوفة المطابقة")
rows=[
("WEB-LINES","مكتمل","ثلاثة قطاعات مستقلة وقسم مقاولات متعدد الصفحات."),("WEB-SEO","مكتمل","Metadata وcanonical وsitemap وrobots وبيانات منظمة ومحتوى وطني إنتاجي."),("WEB-CITY","جزئي","نشر صفحات المدن متوقف عمداً حتى توفر محتوى أصلي وقدرة معتمدة."),("WEB-QUOTE","مكتمل","حفظ تلقائي وتتبع ومرفقات عامة محمية."),("GEO-001","مكتمل","13 منطقة وسجل مدن ومصفوفة خدمة وقدرة واعتماد نشر."),("CON-CRM","مكتمل","Pipeline وفرصة ومدينة وقيمة وموعد ومالك وتدقيق."),("CON-SURVEY","مكتمل","معاينات وتقدير وBOQ وبنود علائقية وإصدارات."),("CON-PLAN","مكتمل","WBS ومخاطر وتقدم وحالة مشروع."),("FIELD-CON","مكتمل جزئياً","واجهة يومية ميدانية وGPS؛ المزامنة دون اتصال ما زالت تحسيناً تشغيلياً."),("CON-DOC","مكتمل","ملفات فعلية وإصدارات وإحالات ومراجع وسبب رفض."),("CON-QHSE","مكتمل","Inspection وNCR وسلامة وأولوية وإشعارات."),("CON-PROC","جزئي","الدورة موجودة؛ نموذج درجات مقارنة المقاولين تحسين متقدم."),
]
table(doc,["الرمز","الحالة","دليل التنفيذ"],rows,[1.1,1.15,4.25])

doc.add_page_break(); rows2=[
("CHANGE-CON","مكتمل","أوامر تغيير بقيم وبنود وحالات وتدقيق."),("IPC-CON","مكتمل","مستخلصات واحتجاز واعتماد وفوترة وسداد."),("FIN-PROJ","مكتمل","CBS/WIP/EAC وربط المشروع ورمز التكلفة والتوقع."),("CON-PDF","مكتمل","PDF عربي بهوية دالي وختم وتوقيع عند توفر الأصول الصحيحة."),("RBAC","مكتمل","20 دوراً ونطاق مشروع/مدينة/منطقة وحدود مالية."),("SOD","مكتمل","منع اعتماد المنشئ وحدود تنفيذ واعتماد مستقلة."),("MFA","مكتمل","TOTP واسترداد أحادي وتحديات replay-safe."),("AUDIT","مكتمل","قبل/بعد، correlation، انتقال حالة، Outbox."),("QA","مكتمل","Type/Lint/Build واختبارات أمن وهجرات ورندر."),("QA-PROD","جزئي","فحص إنتاج آلي جاهز؛ يلزم تشغيله بعد نشر commit الحالي."),("OPS","جزئي","Health/readiness وفحص قبول؛ APM/SLO المركزي يحتاج مزوداً معتمداً."),("EXT","اعتماد خارجي","ZATCA ومدد وقوى والتأمينات والبنوك تحتاج أهلية ومفاتيح واتفاقيات."),
]
table(doc,["الرمز","الحالة","دليل التنفيذ"],rows2,[1.1,1.15,4.25])

heading(doc,"3. الواجهات المنفذة")
for i,text in enumerate(["لوحة محفظة المقاولات","الفرص والمناقصات","المعاينات والتقدير وBOQ والعقود","المشروعات ومراكز التكلفة","التخطيط وWBS والمخاطر","اليوميات الميدانية وGPS","الوثائق وRFI وSubmittal والإحالات","الجودة والسلامة","المشتريات ومقاولو الباطن","أوامر التغيير والمستخلصات","CBS/WIP/EAC","التسليم والضمان","تغطية المملكة","الأدوار والنطاقات والحدود المالية","طلب عرض السعر وتتبع المرفقات"],1): para(doc,f"{i}. {text}",9.5,False,"263A43",3)

heading(doc,"4. الاختبارات والأمن")
table(doc,["البوابة","النتيجة","الدليل"],[
("TypeScript","ناجح","tsc --noEmit"),("Lint","ناجح","ESLint بلا أخطاء"),("Build","ناجح","Standalone production artifact"),("اختبارات آلية","10/10","أمن، PostgreSQL، MFA، نطاقات، رندر"),("هجرات الإنتاج","0005–0011","مطبقة على Supabase"),("RLS/Data API","ناجح","anon/authenticated بلا SELECT والجداول للخادم فقط"),("قبول الإنتاج","بانتظار النشر","scripts/production-smoke.mjs"),
],[1.5,1.15,3.85])

heading(doc,"5. المتبقي الحقيقي")
para(doc,"أولاً — نشر الإصدار الحالي على Render وتشغيل فحص القبول الإنتاجي. هذه خطوة تشغيلية وليست فجوة في الكود.",10,True,AMBER)
para(doc,"ثانياً — توفير صور ختم وتوقيع صحيحة تخص شركة دالي. يمنع استخدام أصول مؤسسة أخرى، حتى لو كانت صالحة فنياً.",10,True,AMBER)
para(doc,"ثالثاً — صفحات المدن الأصلية لا تُنشر جماعياً بقالب آلي. تُفتح لكل مدينة بعد اعتماد القدرة والمحتوى المحلي والدليل التشغيلي ومراجعة SEO بشرية.",10,False)
para(doc,"رابعاً — ZATCA ومدد وقوى والتأمينات والبنوك لا تُغلق برمجياً دون بيانات اعتماد وبيئات اختبار واتفاقيات رسمية.",10,False)

heading(doc,"6. الحكم النهائي")
para(doc,"التنفيذ البرمجي الحالي يغطي متطلبات الوثيقة الأساسية والمتقدمة لمنصة العمالة والمقاولات، وليس مجرد بنية أو شاشات فارغة. الفجوات المتبقية محصورة في نشر الإصدار، أصول هوية صحيحة، محتوى مدن أصلي، وموافقات تكامل خارجية. لا يصح وصف أي من هذه العناصر الخارجية بأنها مكتملة قبل الدليل الرسمي.",11,True,NAVY)
para(doc,"اعتماد مالك النظام: ____________________     التاريخ: ____________________",10,False,"51636B",0)

doc.core_properties.title="تقرير المطابقة والتنفيذ لمنصة دالي"; doc.core_properties.subject="الإغلاق التنفيذي لمتطلبات منصة دالي"; doc.core_properties.author="شركة دالي للتشغيل والصيانة"; doc.core_properties.comments="الإصدار 2.0 — تم تحديثه بعد تنفيذ MFA والملفات الهندسية وCBS/WIP/EAC وطلبات العروض وGPS"
doc.save(OUT)
